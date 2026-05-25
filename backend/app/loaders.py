"""File loaders: PDF, DOCX, XLSX, CSV, TXT.

Each loader returns a list of (text, metadata) tuples. One file may yield
multiple records (e.g. one per spreadsheet row) so structured data can be
retrieved more precisely later.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, List, Tuple

import pandas as pd
from docx import Document
from pypdf import PdfReader

from .logger import logger

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".csv", ".txt"}

Record = Tuple[str, dict]


def _row_to_text(row: pd.Series) -> str:
    parts = []
    for col, val in row.items():
        if pd.isna(val):
            continue
        parts.append(f"{col}: {val}")
    return " | ".join(parts)


def load_pdf(path: Path) -> List[Record]:
    records: List[Record] = []
    try:
        reader = PdfReader(str(path))
        for i, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                records.append((text, {"source": path.name, "type": "pdf", "page": i}))
    except Exception as e:
        logger.error(f"PDF load failed for {path}: {e}")
    return records


def load_docx(path: Path) -> List[Record]:
    records: List[Record] = []
    try:
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        text = "\n".join(paragraphs).strip()
        if text:
            records.append((text, {"source": path.name, "type": "docx"}))
    except Exception as e:
        logger.error(f"DOCX load failed for {path}: {e}")
    return records


def load_xlsx(path: Path) -> List[Record]:
    records: List[Record] = []
    try:
        xls = pd.ExcelFile(path)
        for sheet in xls.sheet_names:
            df = xls.parse(sheet).fillna("")
            for idx, row in df.iterrows():
                text = _row_to_text(row)
                if text:
                    records.append(
                        (
                            text,
                            {
                                "source": path.name,
                                "type": "xlsx",
                                "sheet": sheet,
                                "row": int(idx) + 2,  # +2 to account for header row
                            },
                        )
                    )
    except Exception as e:
        logger.error(f"XLSX load failed for {path}: {e}")
    return records


def load_csv(path: Path) -> List[Record]:
    records: List[Record] = []
    try:
        df = pd.read_csv(path).fillna("")
        for idx, row in df.iterrows():
            text = _row_to_text(row)
            if text:
                records.append(
                    (text, {"source": path.name, "type": "csv", "row": int(idx) + 2})
                )
    except Exception as e:
        logger.error(f"CSV load failed for {path}: {e}")
    return records


def load_txt(path: Path) -> List[Record]:
    try:
        text = path.read_text(encoding="utf-8", errors="ignore").strip()
        if text:
            return [(text, {"source": path.name, "type": "txt"})]
    except Exception as e:
        logger.error(f"TXT load failed for {path}: {e}")
    return []


_LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".xlsx": load_xlsx,
    ".xls": load_xlsx,
    ".csv": load_csv,
    ".txt": load_txt,
}


def load_file(path: Path) -> List[Record]:
    ext = path.suffix.lower()
    loader = _LOADERS.get(ext)
    if not loader:
        logger.warning(f"Unsupported file type: {path}")
        return []
    return loader(path)


def iter_company_files(base_dir: Path) -> Iterable[Path]:
    for p in base_dir.rglob("*"):
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield p
