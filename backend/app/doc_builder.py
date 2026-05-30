"""Universal document builder.

Turns a simple spec into a downloadable file in one of four formats:
    pdf  · docx · xlsx · md (also txt)

The spec shape (produced by the LLM, see chat_service):
    {
      "type": "pdf" | "docx" | "xlsx" | "md" | "txt",
      "filename": "warning_letter",     # no extension
      "title": "Official Warning",      # optional
      "body": "## Heading\n- bullet\n…" # markdown-ish content
      "table": {                        # optional, mainly for xlsx
        "columns": ["Name", "Iqama"],
        "rows": [["Ahmed", "1073..."], ...]
      }
    }

Arabic is fully supported in every format. For PDF we reshape + bidi the
text and bundle Amiri/Noto + DejaVu fonts.
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Dict, List, Optional

from .config import settings
from .logger import logger

FONTS_DIR = Path(__file__).resolve().parent / "fonts"

SUPPORTED_TYPES = {"pdf", "docx", "xlsx", "md", "txt"}
DEFAULT_GROUP = "Documents"

_ARABIC_RE = re.compile(r"[؀-ۿݐ-ݿࢠ-ࣿﭐ-﻿]")


def _has_arabic(text: str) -> bool:
    return bool(_ARABIC_RE.search(text or ""))


def _safe_filename(name: str) -> str:
    name = (name or "document").strip()
    name = re.sub(r"\.[a-zA-Z0-9]{1,5}$", "", name)          # drop any extension
    name = re.sub(r"[^\w؀-ۿ\- ]+", "", name)        # keep word chars + arabic
    name = name.strip().replace(" ", "_")
    return name or "document"


# ----------------------------------------------------------------------------
# Tiny markdown tokenizer — enough for headings, lists, rules, bold, paragraphs
# ----------------------------------------------------------------------------

class Block:
    def __init__(self, kind: str, text: str = "", level: int = 0, ordered: bool = False):
        self.kind = kind          # h1..h3 | li | p | hr | blank
        self.text = text
        self.level = level
        self.ordered = ordered


def _strip_inline(text: str) -> str:
    # remove markdown emphasis markers for plain renderers (docx/pdf apply bold separately)
    return re.sub(r"\*\*(.+?)\*\*", r"\1", re.sub(r"(?<!\*)\*(?!\*)(.+?)\*", r"\1", text))


def _parse_blocks(md: str) -> List[Block]:
    blocks: List[Block] = []
    for raw in (md or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            blocks.append(Block("blank"))
            continue
        if re.match(r"^\s*([-*_])\1{2,}\s*$", line):
            blocks.append(Block("hr"))
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            lvl = min(len(m.group(1)), 3)
            blocks.append(Block(f"h{lvl}", m.group(2).strip()))
            continue
        m = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if m:
            blocks.append(Block("li", m.group(1).strip(), ordered=False))
            continue
        m = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if m:
            blocks.append(Block("li", m.group(1).strip(), ordered=True))
            continue
        blocks.append(Block("p", line.strip()))
    return blocks


# ----------------------------------------------------------------------------
# DOCX
# ----------------------------------------------------------------------------

def _build_docx(spec: Dict, out: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.oxml.ns import qn

    doc = Document()

    def _rtl_para(p, text: str):
        if _has_arabic(text):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pPr = p._p.get_or_add_pPr()
            bidi = pPr.makeelement(qn('w:bidi'), {})
            pPr.append(bidi)

    title = spec.get("title")
    if title:
        h = doc.add_heading(title, level=0)
        _rtl_para(h, title)

    for b in _parse_blocks(spec.get("body", "")):
        if b.kind == "blank":
            continue
        if b.kind == "hr":
            doc.add_paragraph("_" * 30)
            continue
        if b.kind.startswith("h"):
            lvl = int(b.kind[1])
            p = doc.add_heading(_strip_inline(b.text), level=lvl)
            _rtl_para(p, b.text)
            continue
        if b.kind == "li":
            style = "List Number" if b.ordered else "List Bullet"
            p = doc.add_paragraph(_strip_inline(b.text), style=style)
            _rtl_para(p, b.text)
            continue
        # paragraph with bold support
        p = doc.add_paragraph()
        _rtl_para(p, b.text)
        for seg, bold in _bold_segments(b.text):
            run = p.add_run(seg)
            run.bold = bold
            run.font.size = Pt(11)

    # optional table
    table = spec.get("table")
    if table and table.get("columns"):
        cols = table["columns"]
        rows = table.get("rows", [])
        t = doc.add_table(rows=1, cols=len(cols))
        t.style = "Light Grid Accent 1"
        for i, c in enumerate(cols):
            t.rows[0].cells[i].text = str(c)
        for r in rows:
            cells = t.add_row().cells
            for i, val in enumerate(r[:len(cols)]):
                cells[i].text = str(val)

    doc.save(str(out))


def _bold_segments(text: str):
    """Yield (segment, is_bold) tuples splitting on **bold** markers."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            yield part[2:-2], True
        else:
            yield part, False


# ----------------------------------------------------------------------------
# XLSX
# ----------------------------------------------------------------------------

def _build_xlsx(spec: Dict, out: Path) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = (spec.get("title") or "Sheet1")[:31]

    table = spec.get("table") or {}
    cols = table.get("columns")
    rows = table.get("rows", [])

    if not cols:
        # No structured table → dump the body text line by line into column A.
        ws["A1"] = spec.get("title") or "Document"
        ws["A1"].font = Font(bold=True, size=14)
        r = 3
        for b in _parse_blocks(spec.get("body", "")):
            if b.kind == "blank":
                r += 1
                continue
            ws.cell(row=r, column=1, value=_strip_inline(b.text))
            r += 1
        wb.save(str(out))
        return

    header_fill = PatternFill("solid", fgColor="6366F1")
    for i, c in enumerate(cols, start=1):
        cell = ws.cell(row=1, column=i, value=str(c))
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    for ri, row in enumerate(rows, start=2):
        for ci, val in enumerate(row[:len(cols)], start=1):
            ws.cell(row=ri, column=ci, value=val)
    # auto width
    for i, c in enumerate(cols, start=1):
        width = max(len(str(c)), *(len(str(r[i-1])) for r in rows if len(r) >= i)) if rows else len(str(c))
        ws.column_dimensions[chr(64 + i)].width = min(max(width + 2, 10), 50)

    wb.save(str(out))


# ----------------------------------------------------------------------------
# Markdown / TXT
# ----------------------------------------------------------------------------

def _build_md(spec: Dict, out: Path, plain: bool = False) -> None:
    parts: List[str] = []
    title = spec.get("title")
    if title:
        parts.append(title if plain else f"# {title}")
        parts.append("")
    body = spec.get("body", "")
    parts.append(_strip_inline(body) if plain else body)
    table = spec.get("table")
    if table and table.get("columns"):
        cols = table["columns"]
        parts.append("")
        if plain:
            parts.append("\t".join(map(str, cols)))
            for r in table.get("rows", []):
                parts.append("\t".join(map(str, r)))
        else:
            parts.append("| " + " | ".join(map(str, cols)) + " |")
            parts.append("| " + " | ".join("---" for _ in cols) + " |")
            for r in table.get("rows", []):
                parts.append("| " + " | ".join(map(str, r)) + " |")
    out.write_text("\n".join(parts) + "\n", encoding="utf-8")


# ----------------------------------------------------------------------------
# PDF (with Arabic shaping)
# ----------------------------------------------------------------------------

def _shape_ar(text: str) -> str:
    """Reshape + bidi an Arabic (or mixed) string for correct PDF rendering."""
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        return get_display(arabic_reshaper.reshape(text))
    except Exception:
        return text


def _build_pdf(spec: Dict, out: Path) -> None:
    from fpdf import FPDF

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    # Register fonts: Arabic (Noto Naskh) + Latin (Noto Sans). Both are
    # full Unicode TTFs, so even if one is missing we can fall back to the
    # other for ALL text (never to helvetica, which is Latin-1 only and
    # crashes on Arabic / smart punctuation / bullets).
    ar_reg = FONTS_DIR / "NotoNaskhArabic-Regular.ttf"
    ar_bld = FONTS_DIR / "NotoNaskhArabic-Bold.ttf"
    la_reg = FONTS_DIR / "NotoSans-Regular.ttf"
    la_bld = FONTS_DIR / "NotoSans-Bold.ttf"
    have_ar = ar_reg.exists()
    have_la = la_reg.exists()
    if have_ar:
        pdf.add_font("AR", "", str(ar_reg))
        pdf.add_font("AR", "B", str(ar_bld if ar_bld.exists() else ar_reg))
    if have_la:
        pdf.add_font("LA", "", str(la_reg))
        pdf.add_font("LA", "B", str(la_bld if la_bld.exists() else la_reg))

    # Pick the best available Unicode family for a given script, never helvetica.
    def _family(arabic: bool) -> str:
        if arabic:
            if have_ar: return "AR"
            if have_la: return "LA"
        else:
            if have_la: return "LA"
            if have_ar: return "AR"
        return "helvetica"  # last resort (ASCII only)

    def set_font(size: int, bold: bool, arabic: bool):
        fam = _family(arabic)
        try:
            pdf.set_font(fam, "B" if bold else "", size)
        except Exception:
            try:
                pdf.set_font(fam, "", size)
            except Exception:
                pdf.set_font("helvetica", "", size)

    def write_line(text: str, size: int = 11, bold: bool = False, gap: float = 2.0):
        if not text.strip():
            pdf.ln(4)
            return
        arabic = _has_arabic(text)
        set_font(size, bold, arabic)
        display = _shape_ar(text) if arabic else text
        align = "R" if arabic else "L"
        pdf.multi_cell(0, size * 0.55 + 3, display, align=align)
        pdf.ln(gap)

    title = spec.get("title")
    if title:
        write_line(title, size=18, bold=True, gap=4)

    for b in _parse_blocks(spec.get("body", "")):
        if b.kind == "blank":
            pdf.ln(2); continue
        if b.kind == "hr":
            y = pdf.get_y()
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(4); continue
        if b.kind.startswith("h"):
            sz = {1: 16, 2: 14, 3: 12}[int(b.kind[1])]
            write_line(_strip_inline(b.text), size=sz, bold=True, gap=2)
            continue
        if b.kind == "li":
            txt = _strip_inline(b.text)
            # Use a dash as the bullet glyph — the Arabic font lacks U+2022 (•),
            # and a dash renders cleanly in both Noto Naskh and Noto Sans.
            write_line((txt + "  -") if _has_arabic(txt) else ("-  " + txt), size=11)
            continue
        write_line(_strip_inline(b.text), size=11)

    # optional table
    table = spec.get("table")
    if table and table.get("columns"):
        pdf.ln(3)
        cols = table["columns"]
        col_w = (pdf.w - pdf.l_margin - pdf.r_margin) / len(cols)
        set_font(11, True, any(_has_arabic(str(c)) for c in cols))
        for c in cols:
            txt = _shape_ar(str(c)) if _has_arabic(str(c)) else str(c)
            pdf.cell(col_w, 9, txt, border=1, align="C")
        pdf.ln()
        for row in table.get("rows", []):
            for val in row[:len(cols)]:
                s = str(val)
                set_font(10, False, _has_arabic(s))
                pdf.cell(col_w, 8, _shape_ar(s) if _has_arabic(s) else s, border=1, align="C")
            pdf.ln()

    pdf.output(str(out))


# ----------------------------------------------------------------------------
# Public entry
# ----------------------------------------------------------------------------

_BUILDERS = {
    "docx": _build_docx,
    "xlsx": _build_xlsx,
    "pdf":  _build_pdf,
    "md":   lambda s, o: _build_md(s, o, plain=False),
    "txt":  lambda s, o: _build_md(s, o, plain=True),
}


def build_document(spec: Dict, group: str = DEFAULT_GROUP) -> Dict:
    """Render `spec` to a file under generated/<group>/. Returns metadata."""
    doc_type = (spec.get("type") or "pdf").lower().strip()
    if doc_type not in SUPPORTED_TYPES:
        doc_type = "pdf"
    base = _safe_filename(spec.get("filename") or spec.get("title") or "document")
    stamp = date.today().isoformat()
    filename = f"{base}_{stamp}.{doc_type}"

    group_dir = settings.generated_path / group
    group_dir.mkdir(parents=True, exist_ok=True)
    out = group_dir / filename
    # avoid clobber
    if out.exists():
        for i in range(1, 100):
            cand = group_dir / f"{base}_{stamp}_{i}.{doc_type}"
            if not cand.exists():
                out = cand
                filename = cand.name
                break

    builder = _BUILDERS[doc_type]
    builder(spec, out)
    logger.info(f"Built document: {group}/{filename} ({out.stat().st_size} bytes)")
    return {
        "filename": filename,
        "group": group,
        "type": doc_type,
        "size": out.stat().st_size,
        "path": str(out),
    }
