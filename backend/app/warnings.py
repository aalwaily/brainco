"""Warning letter generator.

Loads warning_template.docx from company_data/, replaces placeholders
({{NAME}}, {{IQAMA}}, {{DATE}}, {{REASON}}), and writes a new docx to
generated/warnings/.
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Dict, List, Optional

import pandas as pd
from docx import Document

from .config import settings
from .logger import logger

PLACEHOLDERS = ("NAME", "IQAMA", "DATE", "REASON")

EMPLOYEES_FILENAME = "employees.xlsx"
TEMPLATE_FILENAME = "warning_template.docx"


def _replace_in_paragraph(paragraph, mapping: Dict[str, str]) -> None:
    """Replace {{KEY}} placeholders while preserving paragraph (not run-level)."""
    full = "".join(run.text for run in paragraph.runs)
    if not full:
        return
    new = full
    for k, v in mapping.items():
        new = new.replace("{{" + k + "}}", v)
    if new == full:
        return
    # Replace all text in first run, clear others. Simple but effective for templates.
    if paragraph.runs:
        paragraph.runs[0].text = new
        for r in paragraph.runs[1:]:
            r.text = ""


def _apply_replacements(doc: Document, mapping: Dict[str, str]) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)


def _safe_filename(name: str) -> str:
    keep = re.sub(r"[^A-Za-z0-9_\-؀-ۿ ]+", "", name).strip().replace(" ", "_")
    return keep or "employee"


def find_employee(identifier: str) -> Optional[Dict]:
    """Look up an employee by iqama or name from employees.xlsx."""
    path = settings.company_data_path / EMPLOYEES_FILENAME
    if not path.exists():
        logger.warning(f"{EMPLOYEES_FILENAME} not found at {path}")
        return None
    try:
        df = pd.read_excel(path).fillna("")
    except Exception as e:
        logger.error(f"Failed reading employees.xlsx: {e}")
        return None

    cols = {c.lower().strip(): c for c in df.columns}
    ident = str(identifier).strip()

    def pick(*exact: str, contains: Optional[str] = None) -> Optional[str]:
        for c in exact:
            if c in cols:
                return cols[c]
        if contains:
            for k, c in cols.items():
                if contains in k:
                    return c
        return None

    iqama_col = pick("iqama", contains="iqama")
    name_col  = pick("name", "employee name", "full name", contains="name")
    badge_col = pick("badge", "badge no", "badge id", contains="badge")

    def normalize(row) -> Dict:
        out = {k: ("" if pd.isna(v) else v) for k, v in row.to_dict().items()}
        if name_col:
            out["name"] = out.get(name_col, "")
        if iqama_col:
            out["iqama"] = out.get(iqama_col, "")
        return out

    # 1. exact iqama match
    if iqama_col:
        m = df[df[iqama_col].astype(str).str.strip() == ident]
        if not m.empty:
            return normalize(m.iloc[0])
    # 2. exact badge match
    if badge_col:
        m = df[df[badge_col].astype(str).str.strip().str.lower() == ident.lower()]
        if not m.empty:
            return normalize(m.iloc[0])
    # 3. name contains
    if name_col:
        m = df[df[name_col].astype(str).str.contains(ident, case=False, na=False, regex=False)]
        if not m.empty:
            return normalize(m.iloc[0])
    return None


def generate_warning(
    employee_identifier: str,
    reason: str,
    warning_date: Optional[str] = None,
) -> Dict:
    """Generate a warning .docx.

    Returns dict: {path, filename, employee, missing_placeholders}.
    Raises FileNotFoundError if the template is missing.
    """
    template_path = settings.company_data_path / TEMPLATE_FILENAME
    if not template_path.exists():
        raise FileNotFoundError(
            f"Warning template not found: {template_path}. "
            "Run scripts/create_samples.py to generate a default template."
        )

    employee = find_employee(employee_identifier) or {}
    name = str(employee.get("name") or employee.get("Name") or employee_identifier)
    iqama = str(employee.get("iqama") or employee.get("Iqama") or employee_identifier)
    when = warning_date or date.today().isoformat()

    mapping = {
        "NAME": name,
        "IQAMA": iqama,
        "DATE": when,
        "REASON": reason or "",
    }

    doc = Document(str(template_path))
    _apply_replacements(doc, mapping)

    settings.warnings_path.mkdir(parents=True, exist_ok=True)
    filename = f"warning_{_safe_filename(name)}_{when}.docx"
    out_path = settings.warnings_path / filename
    doc.save(str(out_path))

    # Detect any placeholders that survived (template missing them)
    text_all = "\n".join(p.text for p in Document(str(out_path)).paragraphs)
    missing: List[str] = [p for p in PLACEHOLDERS if "{{" + p + "}}" in text_all]

    logger.info(f"Generated warning: {out_path}")
    return {
        "path": str(out_path),
        "filename": filename,
        "employee": {"name": name, "iqama": iqama},
        "date": when,
        "reason": reason,
        "missing_placeholders": missing,
    }


# ---- Natural-language intent detection ----------------------------------

_WARNING_TRIGGERS = (
    "create warning",
    "generate warning",
    "issue warning",
    "delay warning",
    "warning letter",
    "warning for employee",
    "إنذار",
    "انذار",
    "خطاب إنذار",
)


def detect_warning_intent(message: str) -> Optional[Dict[str, str]]:
    """If the user is asking to generate a warning, extract identifier + reason.

    Returns {"identifier": ..., "reason": ...} or None.
    """
    if not message:
        return None
    low = message.lower()
    if not any(t in low for t in _WARNING_TRIGGERS):
        return None

    # Identifier: look for iqama-like number, or 'employee X'
    ident = None
    m = re.search(r"\b(?:employee|emp|iqama|id)\s*[:#]?\s*([A-Za-z0-9_-]+)", message, re.IGNORECASE)
    if m:
        ident = m.group(1)
    else:
        m = re.search(r"\b(\d{4,})\b", message)
        if m:
            ident = m.group(1)

    # Reason: heuristics. Prefer explicit phrasing, then keyword mapping,
    # then fall back to "for ...". We try because/due to/reason: first,
    # then map common keywords (delay/late/absent), and only use "for X"
    # if X doesn't look like an employee identifier.
    reason = ""
    m = re.search(r"\b(?:because|due to|reason[:\-])\s*(.+)$", message, re.IGNORECASE)
    if m:
        reason = m.group(1).strip().rstrip(".")
    if not reason:
        for kw, label in [
            ("delay", "Repeated delay / lateness to work"),
            ("late", "Repeated lateness to work"),
            ("absent", "Unexcused absence"),
            ("absence", "Unexcused absence"),
            ("misconduct", "Misconduct at the workplace"),
            ("negligence", "Negligence of duties"),
        ]:
            if kw in low:
                reason = label
                break
    if not reason:
        m = re.search(r"\bfor\s+(.+)$", message, re.IGNORECASE)
        if m:
            tail = m.group(1).strip().rstrip(".")
            # ignore if it's just "employee X" / "iqama X" / a bare number
            if not re.match(r"^(employee|emp|iqama|id)\b", tail, re.IGNORECASE) and not re.match(r"^[A-Za-z0-9_-]+$", tail):
                reason = tail
    if not ident:
        return None
    return {"identifier": ident, "reason": reason or "Disciplinary warning"}


def list_generated_warnings() -> List[Dict]:
    out: List[Dict] = []
    base = settings.warnings_path
    if not base.exists():
        return out
    for p in sorted(base.glob("*.docx"), key=lambda x: x.stat().st_mtime, reverse=True):
        st = p.stat()
        out.append(
            {
                "filename": p.name,
                "size": st.st_size,
                "created_at": st.st_mtime,
            }
        )
    return out
