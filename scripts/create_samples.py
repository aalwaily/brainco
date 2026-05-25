"""Create sample files inside company_data/.

Generated:
    employees.xlsx
    warning_template.docx
    contracts.txt
"""
from __future__ import annotations

import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / "company_data"


def make_employees() -> Path:
    path = DATA / "employees.xlsx"
    df = pd.DataFrame(
        [
            {"iqama": "1001", "name": "Ahmed Al-Saud",     "position": "Software Engineer", "email": "ahmed@example.com",   "salary": 12000},
            {"iqama": "1002", "name": "Sara Al-Otaibi",    "position": "Product Manager",   "email": "sara@example.com",    "salary": 15000},
            {"iqama": "1003", "name": "Mohammed Al-Harbi", "position": "Accountant",        "email": "mohammed@example.com", "salary": 9000},
            {"iqama": "1004", "name": "Fatimah Al-Zahrani","position": "HR Specialist",     "email": "fatimah@example.com",  "salary": 10000},
            {"iqama": "1234", "name": "Khalid Al-Dosari",  "position": "Sales Lead",        "email": "khalid@example.com",   "salary": 13500},
            {"iqama": "102",  "name": "Layla Al-Qahtani",  "position": "Designer",          "email": "layla@example.com",    "salary": 11000},
        ]
    )
    df.to_excel(path, index=False)
    return path


def make_template() -> Path:
    path = DATA / "warning_template.docx"
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("OFFICIAL WARNING LETTER")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph("")
    doc.add_paragraph("Date: {{DATE}}")
    doc.add_paragraph("Employee Name: {{NAME}}")
    doc.add_paragraph("Iqama / ID: {{IQAMA}}")
    doc.add_paragraph("")
    doc.add_paragraph("Dear {{NAME}},")
    doc.add_paragraph(
        "This letter serves as a formal warning regarding the following matter: "
        "{{REASON}}."
    )
    doc.add_paragraph(
        "You are kindly requested to address this issue immediately. "
        "Repetition of the same behavior may lead to further disciplinary action "
        "in accordance with the company policy and the Saudi Labor Law."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Human Resources Department")

    doc.save(path)
    return path


def make_contracts() -> Path:
    path = DATA / "contracts.txt"
    path.write_text(
        (
            "Company Policy & Contracts (sample)\n"
            "===================================\n\n"
            "1. Working hours: Sunday to Thursday, 9:00 - 17:00, with one hour break.\n"
            "2. Lateness policy: more than 3 unjustified late arrivals in a month "
            "results in a written warning.\n"
            "3. Annual leave: 21 working days per year for employees with less than "
            "5 years of service, 30 days afterwards.\n"
            "4. Probation period: 90 days, extendable once.\n"
            "5. End-of-service: calculated per the Saudi Labor Law.\n"
            "6. Confidentiality: all company data must remain strictly internal.\n"
        ),
        encoding="utf-8",
    )
    return path


def main() -> int:
    DATA.mkdir(parents=True, exist_ok=True)
    created = [make_employees(), make_template(), make_contracts()]
    for p in created:
        print(f"created: {p}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
